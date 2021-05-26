#stampa fisica magazzino (trasformati e znfd)in cartella new:
#nome files: trasformati:"trs"/znfd:"ppp"
#stampa portafoglio ordini da inizio a fine anno, no filtro trs o ppp, Nome file: ordini - tipo: "ordini a fornitore", causale "ORDFO" - stampa per articolo

# encoding=utf-8
# -*- coding: encoding -*-
import unicodedata
from docx import Document
from docx.shared import Inches
import xlrd
import xlwt
import os
import time
import datetime
import calendar
import sys
import win32api
import shutil
import random
import smtplib
from email.MIMEMultipart import MIMEMultipart
from email.MIMEBase import MIMEBase
from email.MIMEText import MIMEText
from email import Encoders
import smtplib
import time
import datetime
from email.MIMEMultipart import MIMEMultipart
from email.MIMEBase import MIMEBase
from email.MIMEText import MIMEText
from email import Encoders
import os


def ricercaproduttore(lista,codice):
    esito=''
    produttore=[]
    nuovoprodotto=[]
    for elemento in lista:
        riga=elemento.split(';')
        #print riga
        if riga[0]==codice:
            produttore.append(riga[1])
            produttore.append(riga[2])
    if produttore:
        #print produttore
        return produttore
    else:
        #produttore.append('*')
        #produttore.append('*')
        #codice=raw_input(codice+' Prodotto nuovo, inserisci il codice:')
        codiceproduttore=raw_input(codice+' Prodotto nuovo, inserisci il codice produttore:')
        nomeproduttore=raw_input('Prodotto nuovo, inserisci il nome del produttore:')
        f=open('C:\\Users\\f.altarocca\\Desktop\\Scripts\\Lavorazione Dati\\controllo magazzino\\dbproduttori\\lista_prodotti_produttori.csv','a')
        f.write(codice+';'+codiceproduttore+';'+nomeproduttore+'\n')
        f.close()
        produttore.append(codiceproduttore)
        produttore.append(nomeproduttore+'\n')
        return produttore

gmail_user = "felice@biosolidale.it"
gmail_pwd = "BioFelice"
def mail(to, subject, text):
    msg = MIMEMultipart()
    msg['From'] = gmail_user
    msg['To'] = to
    msg['Subject'] = subject
    msg.attach(MIMEText(text))
    mailServer = smtplib.SMTP("smtp.gmail.com", 587)
    mailServer.ehlo()
    mailServer.starttls()
    mailServer.ehlo()
    mailServer.login(gmail_user, gmail_pwd)
    mailServer.sendmail(gmail_user, to, msg.as_string())
    mailServer.close()

def convertXLS2CSV2(aFile,nomefile):
    '''converts a MS Excel file to csv w/ the same name in the same directory'''
    print "------ beginning to convert XLS to CSV ------"
    try:
        import win32com.client, os
        excel = win32com.client.Dispatch('Excel.Application')
        fileDir, fileName = os.path.split(aFile)
        nameOnly = os.path.splitext(fileName)
        newName = "C:\\Users\\f.altarocca\\Desktop\\"+nomefile+".csv"
        outCSV = os.path.join(fileDir, newName)
        workbook = excel.Workbooks.Open(aFile)
        workbook.SaveAs(outCSV, FileFormat=24) # 24 represents xlCSVMSDOS
        workbook.Close(False)
        excel.Quit()
        del excel
        print "...Converted " + nameOnly [0]+ " to CSV"
    except:
        print ">>>>>>> FAILED to convert " + aFile + " to CSV!"
        
def convertXLS2CSV(aFile,nomefile,cartella):
    '''converts a MS Excel file to csv w/ the same name in the same directory'''
    print "------ beginning to convert XLS to CSV ------"
    try:
        import win32com.client, os
        excel = win32com.client.Dispatch('Excel.Application')
        fileDir, fileName = os.path.split(aFile)
        nameOnly = os.path.splitext(fileName)
        newName = "C:\\Users\\f.altarocca\\Desktop\\Scripts\\Lavorazione Dati\\controllo magazzino\\"+cartella+"\\"+nomefile+".csv"
        outCSV = os.path.join(fileDir, newName)
        workbook = excel.Workbooks.Open(aFile)
        workbook.SaveAs(outCSV, FileFormat=24) # 24 represents xlCSVMSDOS
        workbook.Close(False)
        excel.Quit()
        del excel
        print "...Converted " + nameOnly [0]+ " to CSV"
    except:
        print ">>>>>>> FAILED to convert " + aFile + " to CSV!"

def creaLista(nomelista,cartella):
    f=open('C:\\Users\\f.altarocca\\Desktop\\Scripts\\Lavorazione Dati\\controllo magazzino\\dbproduttori\\lista_esclusi.csv','r')
    listaesclusi=f.readlines()
    #print listaesclusi
    f.close()
    nomelista=[]
    files=os.listdir("C:\\Users\\f.altarocca\\Desktop\\Scripts\\Lavorazione Dati\\controllo magazzino\\"+cartella)
    for elemento in files:
        aFile="C:\\Users\\f.altarocca\\Desktop\\Scripts\\Lavorazione Dati\\controllo magazzino\\"+cartella+"\\"+elemento
        nomefile=elemento[:-4]
        #print nomefile
        convertXLS2CSV(aFile,nomefile,cartella)
    files=os.listdir("C:\\Users\\f.altarocca\\Desktop\\Scripts\\Lavorazione Dati\\controllo magazzino\\"+cartella)
    #print files
    for elemento in files:
        if ".csv" in elemento:
            f=open('C:\\Users\\f.altarocca\\Desktop\\Scripts\\Lavorazione Dati\\controllo magazzino\\'+cartella+'\\'+elemento)
            trsold=f.readlines()
            f.close()
            for item in trsold:
                riga=item.split(';')
                codcontrollo=riga[3]+'\n'              
                if riga[3]!='' and riga[3]!='Articolo' and codcontrollo not in listaesclusi:
                    #print riga
                    nomelista.append(riga[3])#codice
                    nomelista.append(riga[4])#descrizione
                    nomelista.append(riga[14])#disponibilita
    return nomelista

def virgolapunto(valore):
    #print valore
    if ',' in valore:
        app=valore.split(',')
        valore='.'.join(app)
        #print type(valore)
    valoreok=float(valore)
    return valoreok

def RicercaCodice(scarti,indnuova,codvecchio,nuova,vecchia):
    indiceok=0
    i=indnuova+1
    #print i
    #print codvecchio
    while i<=len(nuova)-3:
        #print nuova[i]
        if nuova[i]==codvecchio and (i)%3==0:
            indiceok=i
            break
        else:
            i=i+3
    if indiceok==0:
        scarti.append(codvecchio)
    return indiceok

f=open('C:\\Users\\f.altarocca\\Desktop\\Scripts\\lock\\lock.txt','r')
contenuto=f.readlines()
if contenuto[0].upper()=='ON':
    now = datetime.datetime.today()
    anno=str(now)[0:4]
    mese=str(now)[5:7]
    giorno=str(now)[8:10]
    ora=str(now)[11:19]
    app=ora.split(':')
    orario='_'.join(app)
    data=giorno+'_'+mese+'_'+anno+'_'+orario


    #creazione cartelle archivio prec e archiviazione

    MyFile = xlwt.Workbook()
    MyFoglio = MyFile.add_sheet("Situazione Magazzino"+giorno+"_"+mese+"_"+anno, cell_overwrite_ok = True)

    esiste=os.path.exists('C:\\Users\\f.altarocca\\Desktop\\Scripts\\Lavorazione Dati\\controllo magazzino\\prec\\'+anno+'\\')
    if not esiste:
        os.mkdir('C:\\Users\\f.altarocca\\Desktop\\Scripts\\Lavorazione Dati\\controllo magazzino\\prec\\'+anno+'\\')
    esiste=os.path.exists('C:\\Users\\f.altarocca\\Desktop\\Scripts\\Lavorazione Dati\\controllo magazzino\\prec\\'+anno+'\\'+mese+'\\')
    if not esiste:
        os.mkdir('C:\\Users\\f.altarocca\\Desktop\\Scripts\\Lavorazione Dati\\controllo magazzino\\prec\\'+anno+'\\'+mese+'\\')
    esiste=os.path.exists('C:\\Users\\f.altarocca\\Desktop\\Scripts\\Lavorazione Dati\\controllo magazzino\\prec\\'+anno+'\\'+mese+'\\'+giorno+'\\')
    if not esiste:
        os.mkdir('C:\\Users\\f.altarocca\\Desktop\\Scripts\\Lavorazione Dati\\controllo magazzino\\prec\\'+anno+'\\'+mese+'\\'+giorno+'\\')

    #creazione cartelle archivio prec e archiviazione



    document = Document()
               
    vecchia=creaLista('olds','olds')
    nuova=creaLista('new','new')

    for elemento in vecchia:
        if elemento=='':
            indice=vecchia.index(elemento)
            vecchia[indice]='0'
    for elemento in nuova:
        if elemento=='':
            indice=nuova.index(elemento)
            nuova[indice]='0'


    desk=os.listdir('C:\\Users\\f.altarocca\\Desktop\\')
    if 'Stampa_portafoglio_ordini_(Stampa_per_articolo).XLSX' in desk:
        os.rename('C:\\Users\\f.altarocca\\Desktop\\Stampa_portafoglio_ordini_(Stampa_per_articolo).xlsx','C:\\Users\\f.altarocca\\Desktop\\stampa_portafoglio_ordini_(stampa_per_articolo).xls')
    aFile='C:\\Users\\f.altarocca\\Desktop\\stampa_portafoglio_ordini_(stampa_per_articolo).xls'
    nomefile='ordini'
    convertXLS2CSV2(aFile,nomefile)
    f=open('C:\\Users\\f.altarocca\\Desktop\\ordini.csv','r')
    ordini=f.readlines()
    f.close()
    articoliordinati=[]
    for elemento in ordini:
        riga=elemento.split(';')
        riga.pop(len(riga)-1)
        articoliordinati.append(riga[11]) #codice
        articoliordinati.append(riga[10]) #data arrivo
        
    f=open('C:\\Users\\f.altarocca\\Desktop\\Scripts\\Lavorazione Dati\\controllo magazzino\\dbproduttori\\lista_prodotti_produttori.csv','r')
    listaproduttori=f.readlines()
    f.close()

    prodotti_bioso_zero=[]
    scarti=[]
    dazeroauno=[]
    daunoazero=[]
    i=0
    while i<=len(vecchia)-3:
        codvecchio=vecchia[i]
        if codvecchio in nuova:
            indnuova=nuova.index(codvecchio)
            if (indnuova)%3 !=0:
                indnuova=RicercaCodice(scarti,indnuova,codvecchio,nuova,vecchia)
            if indnuova!=0:    
                #print 'Vecchia: '+vecchia[i]+' '+vecchia[i+1]+' '+vecchia[i+2]
                #print 'nuova: '+nuova[indnuova]+' '+nuova[indnuova+1]+' '+nuova[indnuova+2]
                if vecchia[i+2] !=  nuova[indnuova+2]:
                    vecchiovalore=virgolapunto(vecchia[i+2])
                    nuovovalore=virgolapunto(nuova[indnuova+2])
                    if vecchiovalore>0 and nuovovalore==0:
                        daunoazero.append(nuova[indnuova])
                        daunoazero.append(nuova[indnuova+1])
                        daunoazero.append(nuova[indnuova+2])
                        produttore=ricercaproduttore(listaproduttori,nuova[indnuova])
                        #print produttore
                        if produttore[0]=='615':
                            prodottoandatoazero=nuova[indnuova]+' - '+nuova[indnuova+1]
                            prodotti_bioso_zero.append(prodottoandatoazero)
                        daunoazero.append(produttore[0])
                        daunoazero.append(produttore[1])
                    elif vecchiovalore==0 and nuovovalore>0:
                        dazeroauno.append(nuova[indnuova])
                        dazeroauno.append(nuova[indnuova+1])
                        dazeroauno.append(nuova[indnuova+2])
                        produttore=ricercaproduttore(listaproduttori,nuova[indnuova])
                        dazeroauno.append(produttore[0])
                        dazeroauno.append(produttore[1])
        else:
            scarti.append(vecchia[i])
            #scarti.append(vecchia[i+1])
            #scarti.append(vecchia[i+2])
        i=i+3


    #print daunoazero
    #print dazeroauno

    #print articoliordinati

    esiste=os.path.exists('\\\\Dc02\\bioso-car\\ACQUISTI\\SAZ\\'+anno+'\\')
    if not esiste:
        os.mkdir('\\\\Dc02\\bioso-car\\ACQUISTI\\SAZ\\'+anno+'\\')
    esiste=os.path.exists('\\\\Dc02\\bioso-car\\ACQUISTI\\SAZ\\'+anno+'\\'+mese+'\\')
    if not esiste:
        os.mkdir('\\\\Dc02\\bioso-car\\ACQUISTI\\SAZ\\'+anno+'\\'+mese+'\\')
    esiste=os.path.exists('\\\\Dc02\\bioso-car\\ACQUISTI\\SAZ\\'+anno+'\\'+mese+'\\'+giorno+'\\')
    if not esiste:
        os.mkdir('\\\\Dc02\\bioso-car\\ACQUISTI\\SAZ\\'+anno+'\\'+mese+'\\'+giorno+'\\')

    f=open('\\\\Dc02\\bioso-car\\ACQUISTI\\SAZ\\'+anno+'\\'+mese+'\\'+giorno+'\\'+'sottoscorta'+data+'.txt','w')
    f.write('ARTICOLI ENTRATI:\n\n')
    paragraph = document.add_paragraph(style='IntenseQuote')
    paragraph_format = paragraph.paragraph_format
    paragraph_format.left_indent
    None # indicating indentation is inherited from the style hierarchy
    paragraph_format.left_indent = Inches(3)
    paragraph.add_run(data)
    p = document.add_heading('ARTICOLI ENTRATI',0)
    p = document.add_paragraph(style='Normal')
    riga=0
    if dazeroauno:
        MyFoglio.write(riga,0,"PRODOTTO ENTRATI")
        riga=riga+1
        i=0
        while i<=len(dazeroauno)-5:
            if dazeroauno[i+3] !='*':
                riga=riga+1
                #print dazeroauno[i]
                MyFoglio.write(riga,0,dazeroauno[i])
                p.add_run('PRODUTTORE:').bold = True
                p.add_run(dazeroauno[i+3]+' - '+dazeroauno[i+4])
                p.add_run('PRODOTTO:').bold = True
                p.add_run(dazeroauno[i]+' - '+dazeroauno[i+1].decode("iso-8859-1")+': '+dazeroauno[i+2]+' Pz/Kg\n')
                f.write('Produttore: '+dazeroauno[i+3]+' - '+dazeroauno[i+4])
                f.write('**Prodotto: '+dazeroauno[i]+' - '+dazeroauno[i+1]+': '+dazeroauno[i+2]+' Pz/Kg\n')
                if dazeroauno[i] in articoliordinati:
                    indice=articoliordinati.index(dazeroauno[i])
                    dataarrivo=str(articoliordinati[indice+1][:-6])
                    p.add_run('Articolo ordinato.').bold = True
                    p.add_run('In arrivo: '+dataarrivo+'\n\n')
                    f.write('Articolo ordinato. In arrivo: '+dataarrivo+'\n\n')
                else:
                    p.add_run('\n')
            i=i+5
    else:
        MyFoglio.write(riga,0,"NESSUN PRODOTTO RISULTA ENTRATO")
        p.add_run('NESSUN PRODOTTO RISULTA ENTRATO\n')
        f.write('NESSUN PRODOTTO RISULTA ENTRATO\n')
    f.write('\n------------------------------\n\n')
    f.write('ARTICOLI ANDATI A ZERO:\n\n')
    p2 = document.add_heading('ARTICOLI ANDATI A ZERO', 0)
    p2 = document.add_paragraph(style='Normal')
    if daunoazero:
        riga=riga+2
        MyFoglio.write(riga,0,"PRODOTTI ANDATI A ZERO")
        riga=riga+1
        i=0
        while i<=len(daunoazero)-5:
            if daunoazero[i+3] !='*':
                riga=riga+1
                #print dazeroauno[i]
                MyFoglio.write(riga,0,daunoazero[i])
                p2.add_run('PRODUTTORE: ').bold = True
                p2.add_run(daunoazero[i+3]+' - '+daunoazero[i+4])
                p2.add_run('PRODOTTO: ').bold = True
                p2.add_run(daunoazero[i]+' - '+daunoazero[i+1].decode("iso-8859-1")+': '+daunoazero[i+2]+' Pz/Kg\n')
                f.write('Produttore: '+daunoazero[i+3]+' - '+daunoazero[i+4]+'\n')
                f.write('**Prodotto: '+daunoazero[i]+' - '+daunoazero[i+1]+': '+daunoazero[i+2]+' Pz/Kg\n')
                if daunoazero[i] in articoliordinati:
                    indice=articoliordinati.index(daunoazero[i])
                    dataarrivo=str(articoliordinati[indice+1][:-6])
                    p2.add_run('Articolo ordinato. ').bold = True
                    p2.add_run('In arrivo: '+dataarrivo+'\n\n')
                    f.write('Articolo ordinato. In arrivo: '+dataarrivo+'\n\n')
                else:
                    p2.add_run('\n')
            i=i+5
    else:
        MyFoglio.write(riga,0,"NESSUN PRODOTTO ANDATO A ZERO")
        p2.add_run('NESSUN PRODOTTO RISULTA ANDATO A ZERO\n')
        f.write('NESSUN PRODOTTO RISULTA ANDATO A ZERO')

    if scarti:
        f.write('\n\n PRODOTTI NON TROVATI IN STAMPA DI MAGAZZINO NUOVA\n')
        for elemento in scarti:
            f.write(elemento+'\n')
            if elemento in articoliordinati:
                indiceord=articoliordinati.index(elemento)
                f.write('ORDINATO, IN ARRIVO: '+articoliordinati[indiceord]+'\n')
    f.close()


    document.save('\\\\Dc02\\bioso-car\\ACQUISTI\\SAZ\\'+anno+'\\'+mese+'\\'+giorno+'\\'+'sottoscorta'+data+'.doc')


    if prodotti_bioso_zero:
        document_zero=''
        soggetto_zero='Prodotti Biosolidale andati a zero'
        for elemento in prodotti_bioso_zero:
             document_zero=document_zero+elemento+'\n'  
        mail("felice@biosolidale.it",soggetto_zero,document_zero)
        mail("acquisti@biosolidale.it",soggetto_zero,document_zero)

    soggetto='file scorte a zero'
    document='Il file delle scorte a zero e\' disponibile al seguente indirizzo:\\\Serverone\\bioso-car\\ACQUISTI\\SAZ\\'+anno+'\\'+mese+'\\'+giorno
    mail("acquisti@biosolidale.it",soggetto,document)
    mail("magazzino@biosolidale.it",soggetto,document)
    mail("gas@biosolidale.it",soggetto,document)
    mail("felice@biosolidale.it",soggetto,document)

    os.unlink('C:\\Users\\f.altarocca\\Desktop\\ordini.csv')
    shutil.move('C:\\Users\\f.altarocca\\Desktop\\stampa_portafoglio_ordini_(stampa_per_articolo).xls',"C:\\Users\\f.altarocca\\Desktop\\Scripts\\Lavorazione Dati\\controllo magazzino\\prec\\"+anno+'\\'+mese+'\\'+giorno+'\\'+"stampa_portafoglio_ordini_(stampa_per_articolo).xls")
     
    #win32api.ShellExecute(0,"print",'\\\\Dc02\\bioso-car\\ACQUISTI\\SAZ\\'+anno+'\\'+mese+'\\sottoscorta'+data+'.txt',None,".",0)
    #win32api.ShellExecute(0,"print",'\\\\Dc02\\bioso-car\\ACQUISTI\\SAZ\\'+anno+'\\'+mese+'\\sottoscorta'+data+'.doc',None,".",0)
    print '*****'
    print scarti

    #da old a prec
    files=os.listdir("C:\\Users\\f.altarocca\\Desktop\\Scripts\\Lavorazione Dati\\controllo magazzino\\olds")
    for elemento in files:
        if'.csv' in elemento:
            os.unlink("C:\\Users\\f.altarocca\\Desktop\\Scripts\\Lavorazione Dati\\controllo magazzino\\olds\\"+elemento)
        else:
            shutil.move("C:\\Users\\f.altarocca\\Desktop\\Scripts\\Lavorazione Dati\\controllo magazzino\\olds\\"+elemento,"C:\\Users\\f.altarocca\\Desktop\\Scripts\\Lavorazione Dati\\controllo magazzino\\prec\\"+anno+'\\'+mese+'\\'+giorno+'\\'+elemento)

    #trasportare nuovi file in cartella vecchi
    files=os.listdir("C:\\Users\\f.altarocca\\Desktop\\Scripts\\Lavorazione Dati\\controllo magazzino\\new\\")
    for elemento in files:
        if'.csv' in elemento:
            os.unlink("C:\\Users\\f.altarocca\\Desktop\\Scripts\\Lavorazione Dati\\controllo magazzino\\new\\"+elemento)
        else:
            shutil.move("C:\\Users\\f.altarocca\\Desktop\\Scripts\\Lavorazione Dati\\controllo magazzino\\new\\"+elemento,"C:\\Users\\f.altarocca\\Desktop\\Scripts\\Lavorazione Dati\\controllo magazzino\\olds\\"+elemento)

    #MyFile.save('C:\\Users\\f.altarocca\\Desktop\\magazzino_'+giorno+"_"+mese+"_"+anno+'.xls')
    MyFile.save("\\\\Dc02\\bioso-car\\ACQUISTI\\SAZ\\"+anno+"\\"+mese+"\\"+giorno+"\\magazzino_"+giorno+"_"+mese+"_"+anno+"_"+orario+".xls")

